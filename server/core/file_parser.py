"""
文件解析核心服务
支持 PDF、Word、Excel、TXT、MD 等多种文件格式的解析
"""
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from dataclasses import dataclass, field
import logging
import uuid
import json
import os
import tempfile
import shutil
from datetime import datetime
from enum import Enum

logger = logging.getLogger(__name__)


class FileType(Enum):
    """文件类型枚举"""
    PDF = "pdf"
    WORD = "word"
    EXCEL = "excel"
    TXT = "txt"
    MD = "markdown"
    CSV = "csv"
    JSON = "json"
    UNKNOWN = "unknown"


@dataclass
class ParsedFile:
    """解析后的文件"""
    file_id: str
    filename: str
    file_type: FileType
    file_size: int
    content: str
    chunks: List[Dict[str, Any]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())
    chunk_ids: List[str] = field(default_factory=list)
    vector_collection: Optional[str] = None


@dataclass
class FileChunk:
    """文件分块"""
    chunk_id: str
    file_id: str
    content: str
    chunk_index: int
    start_char: int
    end_char: int
    metadata: Dict[str, Any] = field(default_factory=dict)


class FileParser:
    """文件解析�?""
    
    SUPPORTED_EXTENSIONS = {
        '.pdf': FileType.PDF,
        '.docx': FileType.WORD,
        '.doc': FileType.WORD,
        '.xlsx': FileType.EXCEL,
        '.xls': FileType.EXCEL,
        '.txt': FileType.TXT,
        '.md': FileType.MD,
        '.markdown': FileType.MD,
        '.csv': FileType.CSV,
        '.json': FileType.JSON,
    }
    
    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        storage_dir: str = "data/files"
    ):
        """
        初始化文件解析器
        
        Args:
            chunk_size: 分块大小
            chunk_overlap: 分块重叠
            storage_dir: 文件存储目录
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(parents=True, exist_ok=True)
        
        self._metadata_file = self.storage_dir / "files_metadata.json"
        self._files_metadata: Dict[str, Dict[str, Any]] = self._load_metadata()
    
    def _load_metadata(self) -> Dict[str, Dict[str, Any]]:
        """加载文件元数�?""
        if self._metadata_file.exists():
            try:
                with open(self._metadata_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                logger.warning(f"加载文件元数据失�? {e}")
        return {}
    
    def _save_metadata(self):
        """保存文件元数�?""
        try:
            with open(self._metadata_file, 'w', encoding='utf-8') as f:
                json.dump(self._files_metadata, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f"保存文件元数据失�? {e}")
    
    def get_file_type(self, filename: str) -> FileType:
        """获取文件类型"""
        ext = Path(filename).suffix.lower()
        return self.SUPPORTED_EXTENSIONS.get(ext, FileType.UNKNOWN)
    
    def is_supported(self, filename: str) -> bool:
        """检查文件是否支�?""
        return self.get_file_type(filename) != FileType.UNKNOWN
    
    async def parse_file(
        self,
        file_path: str,
        filename: Optional[str] = None,
        custom_metadata: Optional[Dict[str, Any]] = None
    ) -> ParsedFile:
        """
        解析文件
        
        Args:
            file_path: 文件路径
            filename: 文件名（可选，默认从路径提取）
            custom_metadata: 自定义元数据
            
        Returns:
            解析后的文件对象
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存�? {file_path}")
        
        filename = filename or path.name
        file_type = self.get_file_type(filename)
        
        if file_type == FileType.UNKNOWN:
            raise ValueError(f"不支持的文件类型: {filename}")
        
        file_id = f"file_{uuid.uuid4().hex[:12]}"
        file_size = path.stat().st_size
        
        logger.info(f"开始解析文�? {filename}, 类型: {file_type.value}, 大小: {file_size}")
        
        content = await self._extract_content(file_path, file_type)
        
        chunks = self._create_chunks(content, file_id)
        
        metadata = {
            "original_filename": filename,
            "file_type": file_type.value,
            "file_size": file_size,
            "total_chars": len(content),
            "total_chunks": len(chunks),
            "parse_time": datetime.now().isoformat(),
            **(custom_metadata or {})
        }
        
        parsed_file = ParsedFile(
            file_id=file_id,
            filename=filename,
            file_type=file_type,
            file_size=file_size,
            content=content,
            chunks=[{
                "chunk_id": c.chunk_id,
                "content": c.content,
                "chunk_index": c.chunk_index,
                "start_char": c.start_char,
                "end_char": c.end_char,
                "metadata": c.metadata
            } for c in chunks],
            metadata=metadata,
            chunk_ids=[c.chunk_id for c in chunks]
        )
        
        self._files_metadata[file_id] = {
            "file_id": file_id,
            "filename": filename,
            "file_type": file_type.value,
            "file_size": file_size,
            "total_chars": len(content),
            "total_chunks": len(chunks),
            "created_at": parsed_file.created_at,
            "metadata": metadata
        }
        self._save_metadata()
        
        self.save_file_chunks(file_id, parsed_file.chunks)
        
        logger.info(f"文件解析完成: {filename}, 提取 {len(content)} 字符, 分块 {len(chunks)}")
        
        return parsed_file
    
    async def _extract_content(self, file_path: str, file_type: FileType) -> str:
        """提取文件内容"""
        extractors = {
            FileType.PDF: self._extract_pdf,
            FileType.WORD: self._extract_word,
            FileType.EXCEL: self._extract_excel,
            FileType.TXT: self._extract_text,
            FileType.MD: self._extract_text,
            FileType.CSV: self._extract_csv,
            FileType.JSON: self._extract_json,
        }
        
        extractor = extractors.get(file_type)
        if not extractor:
            raise ValueError(f"不支持的文件类型: {file_type}")
        
        return await extractor(file_path)
    
    async def _extract_pdf(self, file_path: str) -> str:
        """提取 PDF 内容"""
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[�?{i+1} 页]\n{page_text}")
            
            return "\n\n".join(text_parts)
        except ImportError:
            logger.warning("pdfplumber 未安装，尝试使用 PyPDF2")
            try:
                from PyPDF2 import PdfReader
                
                reader = PdfReader(file_path)
                text_parts = []
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[�?{i+1} 页]\n{page_text}")
                
                return "\n\n".join(text_parts)
            except ImportError:
                raise ImportError(
                    "PDF 解析需要安�?pdfplumber �?PyPDF2: "
                    "pip install pdfplumber �?pip install PyPDF2"
                )
    
    async def _extract_word(self, file_path: str) -> str:
        """提取 Word 文档内容"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    text_parts.append("[表格]\n" + "\n".join(table_text))
            
            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError(
                "Word 文档解析需要安�?python-docx: pip install python-docx"
            )
    
    async def _extract_excel(self, file_path: str) -> str:
        """提取 Excel 内容"""
        try:
            from openpyxl import load_workbook
            
            wb = load_workbook(file_path, data_only=True)
            
            text_parts = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_text = [f"[工作�? {sheet_name}]"]
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(
                        str(cell) if cell is not None else ""
                        for cell in row
                    )
                    if row_text.strip():
                        sheet_text.append(row_text)
                
                if len(sheet_text) > 1:
                    text_parts.append("\n".join(sheet_text))
            
            return "\n\n".join(text_parts)
        except ImportError:
            raise ImportError(
                "Excel 解析需要安�?openpyxl: pip install openpyxl"
            )
    
    async def _extract_text(self, file_path: str) -> str:
        """提取纯文本内�?""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        with open(file_path, 'rb') as f:
            content = f.read()
            return content.decode('utf-8', errors='ignore')
    
    async def _extract_csv(self, file_path: str) -> str:
        """提取 CSV 内容"""
        import csv
        
        text_parts = []
        encodings = ['utf-8', 'gbk', 'gb2312']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, newline='') as f:
                    reader = csv.reader(f)
                    for row in reader:
                        row_text = " | ".join(row)
                        if row_text.strip():
                            text_parts.append(row_text)
                break
            except UnicodeDecodeError:
                continue
        
        return "\n".join(text_parts)
    
    async def _extract_json(self, file_path: str) -> str:
        """提取 JSON 内容"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        def json_to_text(obj, indent=0):
            """�?JSON 转换为文�?""
            parts = []
            prefix = "  " * indent
            
            if isinstance(obj, dict):
                for key, value in obj.items():
                    if isinstance(value, (dict, list)):
                        parts.append(f"{prefix}{key}:")
                        parts.append(json_to_text(value, indent + 1))
                    else:
                        parts.append(f"{prefix}{key}: {value}")
            elif isinstance(obj, list):
                for i, item in enumerate(obj):
                    if isinstance(item, (dict, list)):
                        parts.append(f"{prefix}[{i}]:")
                        parts.append(json_to_text(item, indent + 1))
                    else:
                        parts.append(f"{prefix}[{i}]: {item}")
            else:
                parts.append(f"{prefix}{obj}")
            
            return "\n".join(parts)
        
        return json_to_text(data)
    
    def _create_chunks(self, content: str, file_id: str) -> List[FileChunk]:
        """创建文本分块"""
        from rag.text_chunker import TextChunker, get_chunker
        
        chunker = get_chunker(self.chunk_size, self.chunk_overlap)
        text_chunks = chunker.chunk(content, metadata={"file_id": file_id})
        
        chunks = []
        for i, tc in enumerate(text_chunks):
            chunk = FileChunk(
                chunk_id=f"chunk_{uuid.uuid4().hex[:8]}",
                file_id=file_id,
                content=tc.content,
                chunk_index=i,
                start_char=tc.start_index,
                end_char=tc.end_index,
                metadata={
                    "file_id": file_id,
                    "chunk_index": i,
                    **tc.metadata
                }
            )
            chunks.append(chunk)
        
        return chunks
    
    def get_file_metadata(self, file_id: str) -> Optional[Dict[str, Any]]:
        """获取文件元数�?""
        return self._files_metadata.get(file_id)
    
    def list_files(self) -> List[Dict[str, Any]]:
        """列出所有文�?""
        return list(self._files_metadata.values())
    
    def delete_file(self, file_id: str) -> bool:
        """删除文件记录"""
        if file_id in self._files_metadata:
            del self._files_metadata[file_id]
            self._save_metadata()
            logger.info(f"已删除文件记�? {file_id}")
            return True
        return False
    
    def get_file_content(self, file_id: str) -> Optional[str]:
        """获取文件内容（从存储中读取）"""
        metadata = self.get_file_metadata(file_id)
        if not metadata:
            return None
        
        content_file = self.storage_dir / f"{file_id}.txt"
        if content_file.exists():
            with open(content_file, 'r', encoding='utf-8') as f:
                return f.read()
        
        return None
    
    def save_file_content(self, file_id: str, content: str):
        """保存文件内容到存�?""
        content_file = self.storage_dir / f"{file_id}.txt"
        with open(content_file, 'w', encoding='utf-8') as f:
            f.write(content)
    
    def save_file_chunks(self, file_id: str, chunks: List[Dict[str, Any]]):
        """保存文件分块信息"""
        chunks_file = self.storage_dir / f"{file_id}_chunks.json"
        with open(chunks_file, 'w', encoding='utf-8') as f:
            json.dump(chunks, f, ensure_ascii=False, indent=2)


class FileVectorService:
    """文件向量化服�?""
    
    def __init__(self, collection_prefix: str = "files"):
        """
        初始化文件向量化服务
        
        Args:
            collection_prefix: 向量集合前缀
        """
        self.collection_prefix = collection_prefix
        self._embedder = None
        self._vector_store = None
    
    def _get_embedder(self):
        """获取嵌入�?""
        if self._embedder is None:
            from rag.embedder import get_embedder
            self._embedder = get_embedder()
        return self._embedder
    
    def _get_vector_store(self):
        """获取向量存储"""
        if self._vector_store is None:
            from rag.vector_store import get_vector_store
            self._vector_store = get_vector_store()
        return self._vector_store
    
    async def index_file(self, parsed_file: ParsedFile) -> Tuple[str, List[str]]:
        """
        将文件内容索引到向量数据�?        
        Args:
            parsed_file: 解析后的文件
            
        Returns:
            (集合名称, 文档ID列表)
        """
        if not parsed_file.chunks:
            logger.warning(f"文件没有内容�? {parsed_file.file_id}")
            return "", []
        
        collection_name = f"{self.collection_prefix}_{parsed_file.file_id}"
        
        embedder = self._get_embedder()
        vector_store = self._get_vector_store()
        
        chunk_contents = [c["content"] for c in parsed_file.chunks]
        chunk_ids = [c["chunk_id"] for c in parsed_file.chunks]
        
        logger.info(f"开始向量化文件: {parsed_file.filename}, {len(chunk_contents)} 个块")
        
        embeddings = embedder.embed(chunk_contents)
        
        metadatas = [
            {
                "file_id": parsed_file.file_id,
                "filename": parsed_file.filename,
                "file_type": parsed_file.file_type.value,
                "chunk_index": c["chunk_index"],
                "start_char": c["start_char"],
                "end_char": c["end_char"],
                **parsed_file.metadata
            }
            for c in parsed_file.chunks
        ]
        
        doc_ids = vector_store.add_documents(
            collection_name=collection_name,
            documents=chunk_contents,
            embeddings=embeddings,
            metadatas=metadatas,
            ids=chunk_ids
        )
        
        logger.info(f"文件向量化完�? {parsed_file.filename}, 集合: {collection_name}")
        
        return collection_name, doc_ids
    
    async def search_file_content(
        self,
        file_id: str,
        query: str,
        top_k: int = 5
    ) -> List[Dict[str, Any]]:
        """
        在文件中搜索内容
        
        Args:
            file_id: 文件ID
            query: 查询文本
            top_k: 返回结果�?            
        Returns:
            搜索结果列表
        """
        collection_name = f"{self.collection_prefix}_{file_id}"
        
        embedder = self._get_embedder()
        vector_store = self._get_vector_store()
        
        return vector_store.search_by_text(
            collection_name=collection_name,
            query_text=query,
            embedder=embedder,
            top_k=top_k
        )
    
    async def search_all_files(
        self,
        query: str,
        top_k: int = 5
    ) -> List[Dict[str, Any]]:
        """
        在所有文件中搜索
        
        Args:
            query: 查询文本
            top_k: 返回结果�?            
        Returns:
            搜索结果列表
        """
        embedder = self._get_embedder()
        vector_store = self._get_vector_store()
        
        query_embedding = embedder.embed_single(query)
        
        collections = vector_store.list_collections()
        file_collections = [c for c in collections if c.startswith(self.collection_prefix)]
        
        all_results = []
        for collection_name in file_collections:
            try:
                results = vector_store.search(
                    collection_name=collection_name,
                    query_embedding=query_embedding,
                    top_k=top_k
                )
                all_results.extend(results)
            except Exception as e:
                logger.warning(f"搜索集合 {collection_name} 失败: {e}")
        
        all_results.sort(key=lambda x: x.get("score", 0), reverse=True)
        
        return all_results[:top_k]
    
    async def delete_file_vectors(self, file_id: str) -> bool:
        """
        删除文件的向量索�?        
        Args:
            file_id: 文件ID
            
        Returns:
            是否成功
        """
        collection_name = f"{self.collection_prefix}_{file_id}"
        
        try:
            vector_store = self._get_vector_store()
            vector_store.delete_collection(collection_name)
            logger.info(f"已删除文件向量集�? {collection_name}")
            return True
        except Exception as e:
            logger.warning(f"删除文件向量集合失败: {e}")
            return False


_file_parser_instance: Optional[FileParser] = None
_file_vector_instance: Optional[FileVectorService] = None


def get_file_parser(
    chunk_size: int = 500,
    chunk_overlap: int = 50,
    storage_dir: str = "data/files"
) -> FileParser:
    """获取文件解析器实�?""
    global _file_parser_instance
    if _file_parser_instance is None:
        _file_parser_instance = FileParser(chunk_size, chunk_overlap, storage_dir)
    return _file_parser_instance


def get_file_vector_service(collection_prefix: str = "files") -> FileVectorService:
    """获取文件向量化服务实�?""
    global _file_vector_instance
    if _file_vector_instance is None:
        _file_vector_instance = FileVectorService(collection_prefix)
    return _file_vector_instance
